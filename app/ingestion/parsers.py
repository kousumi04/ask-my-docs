"""
Document parsers: one function per file type, all normalizing to the
same internal representation so downstream chunking doesn't care what
format the source document was.

Why separate parsers instead of one "smart" auto-parser: PDFs have
pages and often garbled whitespace from column layouts; DOCX has
paragraph/heading structure; Markdown has its own heading syntax we
want to preserve as metadata. Treating them identically loses
information we can otherwise use for citation quality later (e.g.
"page 4" is a much better citation than "chunk 17").
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument


@dataclass
class ParsedDocument:
    source: str          # original filename
    file_type: str       # pdf | docx | md | txt
    full_text: str        # concatenated, cleaned text
    pages: list[str] = field(default_factory=list)  # per-page text; empty list if not paginated


def _clean_text(text: str) -> str:
    """Collapse excessive whitespace introduced by PDF/DOCX extraction quirks."""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def parse_pdf(path: Path) -> ParsedDocument:
    doc = fitz.open(path)
    pages = [_clean_text(page.get_text()) for page in doc]
    doc.close()
    return ParsedDocument(
        source=path.name,
        file_type="pdf",
        full_text="\n\n".join(pages),
        pages=pages,
    )


def parse_docx(path: Path) -> ParsedDocument:
    doc = DocxDocument(str(path))
    # python-docx has no page concept (pagination is a rendering-time
    # detail in Word, not stored in the file) so we treat the whole
    # document as one "page" for metadata purposes.
    paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
    full_text = _clean_text("\n\n".join(paragraphs))
    return ParsedDocument(source=path.name, file_type="docx", full_text=full_text, pages=[])


def parse_markdown(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8")
    return ParsedDocument(source=path.name, file_type="md", full_text=_clean_text(text), pages=[])


def parse_txt(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8")
    return ParsedDocument(source=path.name, file_type="txt", full_text=_clean_text(text), pages=[])


_PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_txt,
}


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix not in _PARSERS:
        raise ValueError(f"Unsupported file type: {suffix} (supported: {list(_PARSERS)})")
    return _PARSERS[suffix](path)